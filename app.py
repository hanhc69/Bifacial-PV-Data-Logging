import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
from datetime import datetime
import os
import hashlib
import matplotlib.pyplot as plt
from fpdf import FPDF
import tempfile
from supabase import create_client

supabase = create_client(
    st.secrets["SUPABASE_URL"],
    st.secrets["SUPABASE_KEY"]
)

# =========================
# AUTHENTICATION
# =========================

SALT = "pv_secure_salt_2026"

def hash_password(password: str) -> str:
    return hashlib.sha256((password + SALT).encode()).hexdigest()

def check_password(password: str) -> bool:
    return hash_password(password) == hash_password("admin123")

def login():
    st.title("🔐 Login")

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("Admin Login")
        password = st.text_input("Enter Admin Password", type="password", key="admin_pass")

        if st.button("Login as Admin"):
            if check_password(password):
                st.session_state.auth = True
                st.session_state.user_role = "admin"
                st.success("Admin access granted")
                st.rerun()
            else:
                st.error("Wrong password")

    with col2:
        st.subheader("Guest Access")
        if st.button("Login as Guest"):
            st.session_state.auth = True
            st.session_state.user_role = "guest"
            st.success("Guest access granted")
            st.rerun()

# =========================
# PLOTTING FUNCTION
# =========================

def fig_to_image_bytes(fig):
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=300, bbox_inches="tight")
    buf.seek(0)
    return buf

def plot_weather_signals(time, temperatures, irradiances, title="Weather Data"):
    fig, ax1 = plt.subplots(figsize=(12, 6))

    for label, temp_values in temperatures.items():
        ax1.plot(time, temp_values, label=label)

    ax1.set_xlabel("Time")
    ax1.set_ylabel("Temperature (°C)")

    ax2 = ax1.twinx()

    for label, irr_values in irradiances.items():
        ax2.plot(time, irr_values, linestyle="--", label=label)

    ax2.set_ylabel("Irradiance (W/m²)")

    plt.title(title)

    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1 + lines2, labels1 + labels2, loc="upper left")

    fig.tight_layout()
    return fig

# =========================
# PREVIEW FUNCTION
# =========================

def preview_report_content(df, report_title, observation):
    """Display a preview of what will be in the report"""
    start_time = f"{df['Date'].iloc[0]} {df['Time'].iloc[0]}" if "Date" in df.columns and "Time" in df.columns else "N/A"
    end_time = f"{df['Date'].iloc[-1]} {df['Time'].iloc[-1]}" if "Date" in df.columns and "Time" in df.columns else "N/A"
    
    st.write("**Report Metadata:**")
    col1, col2 = st.columns(2)
    with col1:
        st.write(f"• Title: {report_title}")
        st.write(f"• Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        st.write(f"• Records: {df.shape[0]}")
    with col2:
        st.write(f"• Start: {start_time}")
        st.write(f"• End: {end_time}")
        st.write(f"• Columns: {df.shape[1]}")
    
    st.write("**Observation Notes:**")
    st.write(observation if observation else "No notes provided")
    
    st.write("**Column Overview:**")
    st.write(", ".join(df.columns))
    
    st.write("**Numeric Summary:**")
    numeric_df = df.select_dtypes(include="number")
    if not numeric_df.empty:
        summary_data = {
            'Column': numeric_df.columns,
            'Mean': [f"{numeric_df[col].mean():.2f}" for col in numeric_df.columns],
            'Min': [f"{numeric_df[col].min():.2f}" for col in numeric_df.columns],
            'Max': [f"{numeric_df[col].max():.2f}" for col in numeric_df.columns],
        }
        st.dataframe(pd.DataFrame(summary_data), use_container_width=True)
    else:
        st.info("No numeric columns found")

# =========================
# WORD REPORT
# =========================

def generate_word_report(df, report_title, observation, fig):
    doc = Document()
    doc.add_heading(report_title, level=1)

    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"

    start_time = f"{df['Date'].iloc[0]} {df['Time'].iloc[0]}" if "Date" in df.columns and "Time" in df.columns else "N/A"
    end_time = f"{df['Date'].iloc[-1]} {df['Time'].iloc[-1]}" if "Date" in df.columns and "Time" in df.columns else "N/A"

    info = [
        ("Generated Date", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("Total Records", str(df.shape[0])),
        ("Total Columns", str(df.shape[1])),
        ("Start Time", start_time),
        ("End Time", end_time),
        ("Observation Notes", observation)
    ]

    for i, (k, v) in enumerate(info):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v

    doc.add_heading("Column Overview", level=2)
    doc.add_paragraph(", ".join(df.columns))

    numeric_df = df.select_dtypes(include="number")

    if not numeric_df.empty:
        doc.add_heading("Numeric Summary", level=2)
        summary = doc.add_table(rows=len(numeric_df.columns) + 1, cols=4)
        summary.style = "Table Grid"

        summary.cell(0, 0).text = "Column"
        summary.cell(0, 1).text = "Mean"
        summary.cell(0, 2).text = "Min"
        summary.cell(0, 3).text = "Max"

        for i, col in enumerate(numeric_df.columns, start=1):
            summary.cell(i, 0).text = col
            summary.cell(i, 1).text = f"{numeric_df[col].mean():.2f}"
            summary.cell(i, 2).text = f"{numeric_df[col].min():.2f}"
            summary.cell(i, 3).text = f"{numeric_df[col].max():.2f}"

    doc.add_heading("Weather Graph", level=2)
    img_stream = fig_to_image_bytes(fig)
    doc.add_picture(img_stream)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# =========================
# PDF REPORT
# =========================

def generate_pdf_report(df, report_title, observation, fig):
    pdf = FPDF()
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, report_title, new_x="LMARGIN", new_y="NEXT")

    start_time = f"{df['Date'].iloc[0]} {df['Time'].iloc[0]}" if "Date" in df.columns and "Time" in df.columns else "N/A"
    end_time = f"{df['Date'].iloc[-1]} {df['Time'].iloc[-1]}" if "Date" in df.columns and "Time" in df.columns else "N/A"

    info = [
        ("Generated Date", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("Total Records", str(df.shape[0])),
        ("Total Columns", str(df.shape[1])),
        ("Start Time", start_time),
        ("End Time", end_time),
        ("Observation Notes", observation)
    ]

    pdf.ln(5)
    pdf.set_font("Helvetica", size=10)

    for key, value in info:
        pdf.cell(50, 8, key, border=1)
        pdf.cell(130, 8, str(value), border=1)
        pdf.ln()

    pdf.ln(5)
    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Column Overview", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", size=10)
    pdf.multi_cell(0, 8, ", ".join(df.columns))

    numeric_df = df.select_dtypes(include="number")

    if not numeric_df.empty:
        pdf.ln(5)
        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 10, "Numeric Summary", new_x="LMARGIN", new_y="NEXT")

        pdf.set_font("Helvetica", "B", 10)
        headers = ["Column", "Mean", "Min", "Max"]

        for header in headers:
            pdf.cell(45, 8, header, border=1)

        pdf.ln()
        pdf.set_font("Helvetica", size=10)

        for col in numeric_df.columns:
            pdf.cell(45, 8, str(col), border=1)
            pdf.cell(45, 8, f"{numeric_df[col].mean():.2f}", border=1)
            pdf.cell(45, 8, f"{numeric_df[col].min():.2f}", border=1)
            pdf.cell(45, 8, f"{numeric_df[col].max():.2f}", border=1)
            pdf.ln()

    pdf.ln(5)
    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Weather Graph", new_x="LMARGIN", new_y="NEXT")

    img_stream = fig_to_image_bytes(fig)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
        tmp.write(img_stream.getvalue())
        temp_image_path = tmp.name

    pdf.image(temp_image_path, w=180)
    return bytes(pdf.output())

# =========================
# MAIN APP
# =========================

if "auth" not in st.session_state:
    st.session_state.auth = False
    st.session_state.user_role = None

if not st.session_state.auth:
    login()
    st.stop()

st.sidebar.subheader("Session")
st.sidebar.write(f"Role: {st.session_state.user_role.capitalize()}")

if st.sidebar.button("🚪 Logout"):
    st.session_state.auth = False
    st.session_state.user_role = None
    st.rerun()

st.title("📊 Bifacial PV Data Logging System")

file = st.file_uploader("Upload CSV", type=["csv"])

report_title = st.text_input("Report Title", "Bifacial PV Performance Report")
observation = st.text_area("Observation Notes")

if file is not None:

    df = pd.read_csv(file)

    st.subheader("📊 Data Preview")
    st.dataframe(df.head(100))

    st.subheader("📌 Dataset Info")
    st.write(f"Rows: {df.shape[0]}")
    st.write(f"Columns: {df.shape[1]}")

    time = df["Time"] if "Time" in df.columns else df.index

    numeric_cols = df.select_dtypes(include="number").columns.tolist()

    st.subheader("📈 Graph Configuration")

    selected_temps = st.multiselect(
        "Select Temperature Columns",
        numeric_cols,
        default=[c for c in numeric_cols if "temp" in c.lower()]
    )

    selected_irradiance = st.multiselect(
        "Select Irradiance Columns",
        numeric_cols,
        default=[c for c in numeric_cols if "irr" in c.lower()]
    )

    temperatures = {col: df[col].tolist() for col in selected_temps}
    irradiances = {col: df[col].tolist() for col in selected_irradiance}

    if selected_temps or selected_irradiance:
        fig = plot_weather_signals(time, temperatures, irradiances)
        st.pyplot(fig)

        st.subheader("📄 Generate Reports")
        
        if st.button("👁️ Preview Report"):
            st.session_state.show_preview = True
        
        if st.session_state.get("show_preview"):
            with st.expander("📋 Report Preview", expanded=True):
                preview_report_content(df, report_title, observation)
            
            st.divider()
            col1, col2 = st.columns(2)
            
            with col1:
                report = generate_word_report(df, report_title, observation, fig)
                st.download_button(
                    label="⬇️ Download Word Report",
                    data=report,
                    file_name="PV_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            with col2:
                report = generate_pdf_report(df, report_title, observation, fig)
                st.download_button(
                    label="⬇️ Download PDF Report",
                    data=report,
                    file_name="PV_Report.pdf",
                    mime="application/pdf"
                )

    if st.button("Test Supabase"):
        supabase.table("pi_commands").update({"command": "hello"}).eq("id", 1).execute()
        st.success("Database updated!")

if st.session_state.user_role == "admin":
    st.divider()
    st.subheader("🔴 Admin Controls")

    if st.button("🔄 Restart Raspberry Pi"):
        os.system("sudo reboot")

    if st.button("⛔ Shutdown Raspberry Pi"):
        os.system("sudo shutdown -h now")
else:
    st.divider()
    st.info("ℹ️ Admin controls are not available in guest mode.")
